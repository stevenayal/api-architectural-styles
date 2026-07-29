"""Verifica el .docx generado: estructura, estilos, columnas y tablas."""
from pathlib import Path

import docx
from docx.oxml.ns import qn

DOC = Path(__file__).resolve().parents[1] / "docs" / "Tarea3-EstilosIntegracionAPIs.docx"

doc = docx.Document(str(DOC))
body = doc.element.body

print("=== ORDEN DEL CUERPO ===")
for i, el in enumerate(body.iterchildren()):
    tag = el.tag.split("}")[1]
    if tag == "p":
        p_pr = el.find(qn("w:pPr"))
        style = None
        sect = ""
        if p_pr is not None:
            s = p_pr.find(qn("w:pStyle"))
            style = s.get(qn("w:val")) if s is not None else None
            sp = p_pr.find(qn("w:sectPr"))
            if sp is not None:
                c = sp.find(qn("w:cols"))
                n = c.get(qn("w:num")) if c is not None else "1"
                sect = f"   <<< SECT cols={n or 1}"
        txt = "".join(t.text or "" for t in el.iter(qn("w:t")))
        print(f"[{i:3}] {str(style):12} {txt[:64]!r}{sect}")
    elif tag == "tbl":
        rows = el.findall(qn("w:tr"))
        cells = rows[0].findall(qn("w:tc")) if rows else []
        print(f"[{i:3}] TABLA {len(rows)}x{len(cells)}")
    elif tag == "sectPr":
        c = el.find(qn("w:cols"))
        print(f"[{i:3}] SECT FINAL cols={c.get(qn('w:num')) if c is not None else 1}")

print(f"\nparrafos={len(doc.paragraphs)} tablas={len(doc.tables)}")
