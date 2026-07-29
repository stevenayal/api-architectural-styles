"""Inspeccion de la plantilla IEEE: estilos, secciones y parrafos de ejemplo."""
import sys
from pathlib import Path

import docx

TEMPLATE = Path(__file__).resolve().parents[1] / "assets" / "ieee-template-transitional.docx"

doc = docx.Document(str(TEMPLATE))

print("=== ESTILOS DE PARRAFO ===")
for s in doc.styles:
    if s.type is not None and str(s.type).startswith("PARAGRAPH"):
        print(f"  {s.name!r}")

print("\n=== ESTILOS DE TABLA ===")
for s in doc.styles:
    if s.type is not None and str(s.type).startswith("TABLE"):
        print(f"  {s.name!r}")

print("\n=== SECCIONES ===")
for i, sec in enumerate(doc.sections):
    cols = sec._sectPr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols"
    )
    num = cols.get(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"
    ) if cols is not None else None
    print(f"  [{i}] start={sec.start_type} cols={num} page={sec.page_width}x{sec.page_height} "
          f"margins L{sec.left_margin} R{sec.right_margin} T{sec.top_margin} B{sec.bottom_margin}")

print("\n=== PARRAFOS DEL CUERPO (primeros 80) ===")
for i, p in enumerate(doc.paragraphs[:80]):
    txt = p.text.strip().replace("\n", " ")
    print(f"  [{i:3}] style={p.style.name!r:28} :: {txt[:90]!r}")

print(f"\nTOTAL parrafos: {len(doc.paragraphs)}  tablas: {len(doc.tables)}")
for i, t in enumerate(doc.tables):
    print(f"  tabla[{i}] style={t.style.name!r} filas={len(t.rows)} cols={len(t.columns)}")
