"""Genera el .docx del articulo sobre la plantilla oficial IEEE (A4, dos columnas).

La plantilla que publica IEEE viene en OOXML *Strict* y python-docx no la abre;
por eso `scripts/convert_template.ps1` la reguarda con Word como Transitional y
este script trabaja sobre `assets/ieee-template-transitional.docx`.

Estrategia: no se reconstruye el formato, se reutiliza el de la plantilla.
Del cuerpo original solo se conservan los parrafos que llevan `sectPr` (los saltos
de seccion que definen 1 columna para el titulo, N columnas para los autores y
2 columnas para el cuerpo); el contenido de ejemplo se borra y el propio se
inserta antes del parrafo ancla que cierra la seccion a dos columnas.
"""
from __future__ import annotations

import copy
import re
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "assets" / "ieee-template-transitional.docx"
SOURCE = ROOT / "docs" / "articulo.md"
OUTPUT = ROOT / "docs" / "Tarea3-EstilosIntegracionAPIs.docx"

# Indices de la plantilla (ver scripts/inspect_body.py).
P_TITLE = 0
P_AUTHOR_FIRST, P_AUTHOR_LAST = 4, 7
P_AUTHOR_SECT = 8          # salto de seccion de la franja de autores (cols=3)
P_EXTRA_AUTHOR_SECT = 9    # segunda franja de autores: sobra con dos autores
P_BODY_FIRST, P_BODY_LAST = 10, 89
P_ANCHOR = 90              # parrafo que porta el sectPr de 2 columnas del cuerpo
P_TRAILING = 91            # nota final de la plantilla, en seccion de 1 columna


# --------------------------------------------------------------------------- #
# Lectura del articulo
# --------------------------------------------------------------------------- #
def parse(text: str) -> dict:
    doc: dict = {"title": "", "authors": [], "abstract": "", "keywords": "",
                 "blocks": [], "references": []}
    in_refs = False
    pending_caption: str | None = None
    rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal pending_caption, rows
        if pending_caption is not None:
            doc["blocks"].append(("table", pending_caption, rows))
        pending_caption, rows = None, []

    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("%references"):
            flush_table()
            in_refs = True
            continue
        if in_refs:
            doc["references"].append(line)
            continue
        if line.startswith("|"):
            rows.append([c.strip() for c in line.strip("|").split("|")])
            continue
        flush_table()
        if line.startswith("%title:"):
            doc["title"] = line[len("%title:"):].strip()
        elif line.startswith("%author:"):
            doc["authors"].append([p.strip() for p in line[len("%author:"):].split("|")])
        elif line.startswith("%abstract:"):
            doc["abstract"] = line[len("%abstract:"):].strip()
        elif line.startswith("%keywords:"):
            doc["keywords"] = line[len("%keywords:"):].strip()
        elif line.startswith("%tablehead:"):
            # el tabulador del md no sobrevive como texto en OOXML: se normaliza
            pending_caption = " ".join(line[len("%tablehead:"):].split())
        elif line.startswith("## "):
            doc["blocks"].append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            doc["blocks"].append(("h1", line[2:].strip()))
        else:
            doc["blocks"].append(("p", line))
    flush_table()
    return doc


# --------------------------------------------------------------------------- #
# Utilidades de OOXML
# --------------------------------------------------------------------------- #
EMPHASIS = re.compile(r"[*`]([^*`]+)[*`]")


def add_rich_text(paragraph, text: str) -> None:
    """Escribe el texto aplicando cursiva a los tramos entre * o `."""
    pos = 0
    for m in EMPHASIS.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_columns(sect_pr, num: int) -> None:
    cols = sect_pr.find(qn("w:cols"))
    if num <= 1:
        cols.attrib.pop(qn("w:num"), None)
    else:
        cols.set(qn("w:num"), str(num))


def clear_paragraph(p_el) -> None:
    """Borra runs y contenido, preservando pPr (y con el, cualquier sectPr)."""
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)


def set_style(p_el, style_id: str) -> None:
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_el.insert(0, p_pr)
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.insert(0, p_style)
    p_style.set(qn("w:val"), style_id)


def disable_numbering(paragraph) -> None:
    """Anula la numeracion heredada del estilo (numId=0 la desactiva)."""
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = 0


def clear_footers(document) -> None:
    """Vacia todos los pies de pagina (primera pagina, pares e impares)."""
    for section in document.sections:
        for footer in (section.first_page_footer, section.even_page_footer, section.footer):
            footer.is_linked_to_previous = False
            for paragraph in footer.paragraphs:
                for child in list(paragraph._p):
                    if child.tag != qn("w:pPr"):
                        paragraph._p.remove(child)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


# --------------------------------------------------------------------------- #
# Construccion
# --------------------------------------------------------------------------- #
class Builder:
    """Inserta contenido justo antes del parrafo ancla de la seccion a 2 columnas."""

    def __init__(self, document, anchor_el, body_sect_pr):
        self.doc = document
        self.anchor = anchor_el
        self.body_sect_pr = body_sect_pr

    def _place(self, element) -> None:
        self.anchor.addprevious(element)

    def paragraph(self, style: str, text: str = "", rich: bool = True):
        p = self.doc.add_paragraph(style=style)
        if text:
            add_rich_text(p, text) if rich else p.add_run(text)
        self._place(p._p)
        return p

    def section_break(self, columns: int) -> None:
        """Parrafo vacio que porta un sectPr: cambia el numero de columnas."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # marca de parrafo diminuta: el salto de seccion no debe abrir un hueco
        run_pr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "2")
        run_pr.append(sz)
        p._p.get_or_add_pPr().append(run_pr)
        sect_pr = copy.deepcopy(self.body_sect_pr)
        for attr in list(sect_pr.attrib):
            if attr.endswith("}rsidR") or attr.endswith("}rsidRPr"):
                sect_pr.attrib.pop(attr)
        set_columns(sect_pr, columns)
        p._p.get_or_add_pPr().append(sect_pr)
        self._place(p._p)

    def table(self, caption: str, rows: list[list[str]]) -> None:
        """Tabla a ancho de pagina: se aisla en una seccion de una sola columna."""
        self.section_break(2)          # cierra el tramo a dos columnas
        # el estilo 'table head' de la plantilla autonumera en ingles ("TABLE I.");
        # se anula para que el rotulo en espanol del md no salga duplicado
        disable_numbering(self.paragraph("table head", caption, rich=False))
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        set_table_borders(table)
        table.autofit = True
        for r, row in enumerate(rows):
            for c, cell_text in enumerate(row):
                cell = table.cell(r, c)
                cell.text = ""
                p = cell.paragraphs[0]
                set_style(p._p, "tablecolhead" if r == 0 else "tablecopy")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
                add_rich_text(p, cell_text)
        self._place(table._tbl)
        self.section_break(1)          # la tabla queda a ancho completo


def build() -> None:
    document = docx.Document(str(TEMPLATE))
    article = parse(SOURCE.read_text(encoding="utf-8"))
    body = document.element.body
    children = list(body.iterchildren())

    anchor = children[P_ANCHOR]
    body_sect_pr = anchor.find(qn("w:pPr")).find(qn("w:sectPr"))

    # --- bloque de titulo -------------------------------------------------- #
    title_el = children[P_TITLE]
    clear_paragraph(title_el)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = article["title"]
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    title_el.append(run)
    for idx in (1, 2):                       # nota "Sub-titles are not captured..."
        body.remove(children[idx])

    # --- autores ----------------------------------------------------------- #
    authors = article["authors"]
    slots = list(range(P_AUTHOR_FIRST, P_AUTHOR_LAST + 1))
    for slot, author in zip(slots, authors):
        el = children[slot]
        clear_paragraph(el)
        run = OxmlElement("w:r")
        for i, line in enumerate(author):
            if i:
                run.append(OxmlElement("w:br"))
            t = OxmlElement("w:t")
            t.text = line
            t.set(qn("xml:space"), "preserve")
            run.append(t)
        el.append(run)
    for slot in slots[len(authors):]:
        body.remove(children[slot])
    set_columns(children[P_AUTHOR_SECT].find(qn("w:pPr")).find(qn("w:sectPr")),
                max(len(authors), 1))
    body.remove(children[P_EXTRA_AUTHOR_SECT])

    # --- pie de pagina ------------------------------------------------------ #
    # La plantilla trae el marcador de copyright de conferencia
    # ("XXX-X-XXXX-XXXX-X/XX/$XX.00 (c)20XX IEEE") en el pie de la primera
    # pagina. Es un dato que asigna la conferencia; en un trabajo academico se
    # elimina.
    clear_footers(document)

    # --- se vacia el cuerpo de ejemplo ------------------------------------- #
    for el in children[P_BODY_FIRST:P_BODY_LAST + 1]:
        body.remove(el)
    clear_paragraph(anchor)
    set_style(anchor, "Normal")
    clear_paragraph(children[P_TRAILING])

    # --- contenido propio --------------------------------------------------- #
    b = Builder(document, anchor, body_sect_pr)

    p = b.paragraph("Abstract")
    p.add_run("Resumen—").bold = True
    add_rich_text(p, article["abstract"])

    p = b.paragraph("Keywords")
    p.add_run("Palabras clave—").bold = True
    add_rich_text(p, article["keywords"])

    for block in article["blocks"]:
        kind = block[0]
        if kind == "h1":
            b.paragraph("Heading 1", block[1], rich=False)
        elif kind == "h2":
            b.paragraph("Heading 2", block[1], rich=False)
        elif kind == "p":
            b.paragraph("Body Text", block[1])
        elif kind == "table":
            b.table(block[1], block[2])

    b.paragraph("Heading 5", "Referencias", rich=False)
    for entry in article["references"]:
        ref = b.paragraph("references", entry, rich=False)
        disable_numbering(ref)
        ref.paragraph_format.first_line_indent = Pt(-14)
        ref.paragraph_format.left_indent = Pt(14)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT))
    print(f"OK -> {OUTPUT}")
    print(f"   parrafos={len(document.paragraphs)} tablas={len(document.tables)} "
          f"referencias={len(article['references'])}")


if __name__ == "__main__":
    sys.exit(build())
